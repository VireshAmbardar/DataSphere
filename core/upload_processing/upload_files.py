# file_utils.py
import os
import shutil
from uuid import UUID
import streamlit as st
from core.global_settings import UPLOAD_FILE_DIR
from llama_index.core import SimpleDirectoryReader
from docx import Document as DocxDocument
from llama_index.core import Document

os.makedirs(UPLOAD_FILE_DIR, exist_ok=True)

if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = []

def process_uploaded_file(uploaded_file):
    """Save file, process pdf/docx/csv, and return LlamaIndex Document(s)."""
    save_path = os.path.join(UPLOAD_FILE_DIR, f"{uploaded_file.file_id}_{uploaded_file.name}")
    with open(save_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    # save path in state
    st.session_state.uploaded_files.append(save_path)
    mime_type = uploaded_file.type


    # ----- PDF -----
    if mime_type == "application/pdf":
        from llama_index.readers.file import PDFReader
        documents = PDFReader().load_data(save_path)

    # ----- DOCX -----
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = DocxDocument(save_path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        documents = Document(text=text, metadata={"source": save_path})

    # ----- CSV -----
    elif mime_type == "text/csv":
        df = pd.read_csv(save_path)
        text = df.to_csv(index=False)
        documents = Document(text=text, metadata={"source": save_path})

    return "Pass", "pass"


def cleanup_uploaded_files():
    """Remove session-tracked files."""
    for path in st.session_state.get("uploaded_files", []):
        if os.path.exists(path):
            os.remove(path)
    st.session_state.uploaded_files.clear()
